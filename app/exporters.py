"""Research-note exporters: Markdown, HTML, DOCX, PDF.

A note is the dict returned by ``notes_store.get_note`` (or any mapping with
``markdown``, ``candidates``, ``asof``, ``provider``). Every exporter degrades
gracefully when the markdown body is empty (e.g. a note generated with no AI
key): the candidate table is still rendered so the export is never empty and
never raises.

PDF uses xhtml2pdf (reportlab-based, pip-only — no system cairo/pango libs, so
it is safe on Railway Nixpacks). HTML/PDF share one styled template; DOCX is
built natively with python-docx. Markdown body -> HTML reuses the ``markdown``
lib already in requirements.
"""
from __future__ import annotations

import io
from html import escape
from typing import Any, Dict, List, Optional

import markdown as _md

# Institutional, print-friendly CSS embedded into the standalone HTML/PDF.
_CSS = """
@page { size: A4; margin: 1.6cm; }
body { font-family: Georgia, 'Times New Roman', serif; color: #1a1a1a;
       font-size: 11pt; line-height: 1.45; }
h1 { font-size: 19pt; margin: 0 0 2px 0; border-bottom: 2px solid #1a1a1a;
     padding-bottom: 6px; font-family: Helvetica, Arial, sans-serif; }
h2 { font-size: 14pt; margin: 18px 0 6px 0; color: #11365c;
     font-family: Helvetica, Arial, sans-serif; }
h3 { font-size: 12pt; margin: 14px 0 4px 0; color: #11365c;
     font-family: Helvetica, Arial, sans-serif; }
.asof { color: #555; font-size: 10pt; margin: 4px 0 14px 0;
        font-family: Helvetica, Arial, sans-serif; }
.meta { color: #777; font-size: 8.5pt; margin-top: 26px;
        border-top: 1px solid #ccc; padding-top: 6px;
        font-family: Helvetica, Arial, sans-serif; }
table { border-collapse: collapse; width: 100%; margin: 8px 0 14px 0;
        font-size: 9.5pt; font-family: Helvetica, Arial, sans-serif; }
th, td { border: 1px solid #bbb; padding: 4px 6px; text-align: left; }
th { background: #11365c; color: #fff; }
tr:nth-child(even) td { background: #f2f5f9; }
.long { color: #0a6b2e; font-weight: bold; }
.short { color: #9b1c1c; font-weight: bold; }
.reject { color: #9b1c1c; font-weight: bold; }
code, pre { font-family: 'Courier New', monospace; }
"""

_VERDICT_LABEL = {
    "MECHANICAL_DISLOCATION": "Mechanical",
    "BROKEN_STORY": "Broken story (REJECT)",
    "NEEDS_DATA": "Needs data (REJECT)",
}


def _fmt(v: Any, spec: str = "") -> str:
    if v is None:
        return "—"
    try:
        if spec:
            return format(float(v), spec)
        return str(v)
    except Exception:
        return str(v)


def _candidate_rows(candidates: List[Dict[str, Any]]) -> List[List[str]]:
    rows: List[List[str]] = []
    for c in candidates or []:
        side = (c.get("side") or "").lower()
        score = c.get("reversion_score") if side == "long" else c.get("fade_score")
        rows.append([
            str(c.get("ticker") or ""),
            str(c.get("name") or ""),
            "LONG" if side == "long" else "SHORT",
            str(c.get("sector") or ""),
            _fmt(c.get("rank_z"), ".2f"),
            _fmt(c.get("peer_relative_z"), ".2f"),
            _fmt(c.get("rsi"), ".1f"),
            _fmt(score, ".3f"),
            "Idiosyncratic" if c.get("dislocation_type") == "IDIOSYNCRATIC" else "Sector/macro",
            _VERDICT_LABEL.get(c.get("verdict") or "", "—"),
        ])
    return rows


_CAND_HEADERS = ["Ticker", "Name", "Side", "Sector", "Rank z", "Peer-rel z",
                 "RSI", "Score", "Type", "Triage"]


def filename(note: Dict[str, Any], ext: str) -> str:
    asof = str(note.get("asof") or "unknown")
    safe = asof.replace(" ", "_").replace(":", "").replace("/", "-")
    return f"research_note_{safe}.{ext}"


# ---------------------------------------------------------------------------
# Markdown
# ---------------------------------------------------------------------------
def to_markdown(note: Dict[str, Any]) -> str:
    asof = note.get("asof") or "unknown"
    parts: List[str] = [f"# Research Note", f"_As of {asof}_", ""]
    cands = note.get("candidates") or []
    if cands:
        parts.append("## Selected candidates")
        parts.append("")
        parts.append("| " + " | ".join(_CAND_HEADERS) + " |")
        parts.append("| " + " | ".join(["---"] * len(_CAND_HEADERS)) + " |")
        for r in _candidate_rows(cands):
            parts.append("| " + " | ".join(r) + " |")
        parts.append("")
    body = (note.get("markdown") or "").strip()
    if body:
        parts.append("## Analyst note")
        parts.append("")
        parts.append(body)
    elif not cands:
        parts.append("_No candidates selected and no written note available._")
    prov = note.get("provider")
    if prov:
        parts.append("")
        parts.append(f"_Generated via {prov} — qualitative only; signals computed in-app._")
    return "\n".join(parts) + "\n"


# ---------------------------------------------------------------------------
# HTML
# ---------------------------------------------------------------------------
def _candidate_table_html(candidates: List[Dict[str, Any]]) -> str:
    if not candidates:
        return "<p><em>No candidates selected under current settings.</em></p>"
    head = "".join(f"<th>{escape(h)}</th>" for h in _CAND_HEADERS)
    body = []
    for r in _candidate_rows(candidates):
        cells = []
        for i, val in enumerate(r):
            cls = ""
            if i == 2:
                cls = ' class="long"' if val == "LONG" else ' class="short"'
            if i == 9 and "REJECT" in val:
                cls = ' class="reject"'
            cells.append(f"<td{cls}>{escape(val)}</td>")
        body.append("<tr>" + "".join(cells) + "</tr>")
    return (f"<table><thead><tr>{head}</tr></thead>"
            f"<tbody>{''.join(body)}</tbody></table>")


def to_html(note: Dict[str, Any]) -> str:
    asof = escape(str(note.get("asof") or "unknown"))
    cands = note.get("candidates") or []
    body_md = (note.get("markdown") or "").strip()
    body_html = _md.markdown(body_md, extensions=["tables"]) if body_md else ""
    sections = [
        "<h2>Selected candidates</h2>",
        _candidate_table_html(cands),
    ]
    if body_html:
        sections.append("<h2>Analyst note</h2>")
        sections.append(f'<div class="note-body">{body_html}</div>')
    elif not cands:
        sections.append("<p><em>No candidates selected and no written note "
                        "available.</em></p>")
    prov = note.get("provider")
    meta = (f'<div class="meta">Generated via {escape(str(prov))} — qualitative '
            f"only; signals computed in-app. Educational tool, not investment "
            f"advice.</div>") if prov else ('<div class="meta">Educational tool, '
            "not investment advice.</div>")
    return (
        "<!DOCTYPE html><html><head><meta charset='utf-8'>"
        "<title>Research Note</title>"
        f"<style>{_CSS}</style></head><body>"
        "<h1>Research Note</h1>"
        f'<div class="asof">As of {asof}</div>'
        + "".join(sections)
        + meta
        + "</body></html>"
    )


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------
def to_docx_bytes(note: Dict[str, Any]) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()
    title = doc.add_heading("Research Note", level=0)
    asof = str(note.get("asof") or "unknown")
    sub = doc.add_paragraph()
    run = sub.add_run(f"As of {asof}")
    run.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    cands = note.get("candidates") or []
    doc.add_heading("Selected candidates", level=1)
    if cands:
        rows = _candidate_rows(cands)
        table = doc.add_table(rows=1, cols=len(_CAND_HEADERS))
        table.style = "Light Grid Accent 1"
        for i, h in enumerate(_CAND_HEADERS):
            cell = table.rows[0].cells[i]
            cell.text = h
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True
        for row in rows:
            cells = table.add_row().cells
            for i, val in enumerate(row):
                cells[i].text = val
    else:
        doc.add_paragraph("No candidates selected under current settings.")

    body_md = (note.get("markdown") or "").strip()
    if body_md:
        doc.add_heading("Analyst note", level=1)
        _docx_render_markdown(doc, body_md)

    prov = note.get("provider")
    foot = doc.add_paragraph()
    frun = foot.add_run(
        (f"Generated via {prov} — qualitative only; signals computed in-app. "
         if prov else "")
        + "Educational tool, not investment advice."
    )
    frun.italic = True
    frun.font.size = Pt(8)
    frun.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


def _docx_render_markdown(doc, text: str) -> None:
    """Light markdown -> Word mapping: headers, bold-label lines, list items,
    paragraphs. Good enough for the structured note the generator emits."""
    from docx.shared import Pt

    for raw in text.splitlines():
        line = raw.rstrip()
        if not line.strip():
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=2)
        elif line.lstrip().startswith(("- ", "* ")):
            p = doc.add_paragraph(style="List Bullet")
            _docx_add_inline(p, line.lstrip()[2:].strip())
        else:
            p = doc.add_paragraph()
            _docx_add_inline(p, line.strip())


def _docx_add_inline(paragraph, text: str) -> None:
    """Render **bold** spans within a line; everything else is plain text."""
    import re

    pos = 0
    for m in re.finditer(r"\*\*(.+?)\*\*", text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        run = paragraph.add_run(m.group(1))
        run.bold = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


# ---------------------------------------------------------------------------
# PDF (xhtml2pdf / reportlab — no system libs)
# ---------------------------------------------------------------------------
def to_pdf_bytes(note: Dict[str, Any]) -> bytes:
    from xhtml2pdf import pisa

    html = to_html(note)
    bio = io.BytesIO()
    pisa.CreatePDF(src=html, dest=bio, encoding="utf-8")
    return bio.getvalue()


# ---------------------------------------------------------------------------
# Dispatch
# ---------------------------------------------------------------------------
_CONTENT_TYPES = {
    "md": "text/markdown; charset=utf-8",
    "html": "text/html; charset=utf-8",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "pdf": "application/pdf",
}


def export(note: Dict[str, Any], fmt: str) -> tuple[bytes, str, str]:
    """Return (content_bytes, content_type, filename) for the requested format.

    Raises ValueError for an unknown format. Never crashes on empty markdown.
    """
    fmt = (fmt or "md").lower()
    if fmt not in _CONTENT_TYPES:
        raise ValueError(f"unknown export format: {fmt}")
    if fmt == "md":
        data = to_markdown(note).encode("utf-8")
    elif fmt == "html":
        data = to_html(note).encode("utf-8")
    elif fmt == "docx":
        data = to_docx_bytes(note)
    else:  # pdf
        data = to_pdf_bytes(note)
    return data, _CONTENT_TYPES[fmt], filename(note, fmt)
